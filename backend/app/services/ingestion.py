import os
import logging
import time
import traceback
import threading
from typing import List, Dict, Any, Optional
from pathlib import Path
from sqlalchemy.orm import Session
from sqlalchemy import text as sa_text

import docx  # python-docx
import pandas as pd
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import io

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

import re
import json
import uuid
from app.services.gemini_embedding_service import GeminiEmbeddingService
from app.services.generation_service import GenerationService
from app.services.graph_extraction_service import GraphExtractionService
from app.models.document import DocumentChunk, PaperSummary
from app.config import settings

logger = logging.getLogger(__name__)

# ── Phase 5: Global semaphore — max 3 concurrent embedding calls ────────────
_embedding_semaphore = threading.Semaphore(3)

MIN_TEXT_LENGTH = 1000  # Phase 3: minimum characters after PDF parse


def _log_ingestion_error(db: Session, filename: str, stage: str, exc: Exception):
    """Persist a structured ingestion error to the ingestion_errors table."""
    tb = traceback.format_exc()
    try:
        db.execute(
            sa_text("""
                INSERT INTO ingestion_errors (filename, stage, error_message, stack_trace)
                VALUES (:fn, :st, :em, :tb)
            """),
            {"fn": filename, "st": stage, "em": str(exc), "tb": tb}
        )
        db.commit()
    except Exception as log_err:
        # Never let error logging crash the caller
        logger.error(f"[ERROR_LOG FAIL] Could not persist error for {filename}: {log_err}")
        try:
            db.rollback()
        except Exception:
            pass


class DocumentIngestionService:
    """
    Service for ingesting documents into the RAG system.

    Pipeline:
    1. Load documents (PDF, DOCX, TXT, MD, CSV)
    2. Validate parsed text length
    3. Split into chunks
    4. Generate embeddings (semaphore-limited, retry-capable)
    5. Store chunks individually (per-chunk safety)
    """

    def __init__(self, db: Session):
        self.db = db
        self.embeddings = GeminiEmbeddingService()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

        # Init Neo4j Driver optionally for GraphRAG
        self.graph_extractor = None
        try:
            from neo4j import GraphDatabase
            driver = GraphDatabase.driver(
                settings.NEO4J_URI,
                auth=(settings.NEO4J_USER, settings.NEO4J_PASSWORD)
            )
            self.graph_extractor = GraphExtractionService(driver)
            logger.info("GraphExtractionService initialized successfully.")
        except Exception as e:
            logger.warning(f"Neo4j driver not initialized. Graph RAG disabled during ingestion: {e}")

    def _perform_ocr(self, file_path: str) -> str:
        """Perform OCR on a PDF file using Tesseract."""
        try:
            logger.info(f"Starting OCR for {file_path}")
            images = convert_from_path(file_path)
            full_text = ""
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                full_text += text + "\n"
            return full_text
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            raise ValueError(f"OCR processing failed: {str(e)}")

    def load_document(self, file_path: str, file_type: str) -> List[Document]:
        """Load a document using appropriate loader."""
        if file_type == "pdf":
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            total_text_len = sum(len(d.page_content) for d in docs)
            if total_text_len < 200:
                logger.info(f"PDF text length {total_text_len} is too short. Triggering OCR fallback.")
                ocr_text = self._perform_ocr(file_path)
                return [Document(page_content=ocr_text, metadata={"source": file_path, "file_type": "pdf", "ingestion_method": "ocr"})]
            return docs

        elif file_type in ["txt", "md"]:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    text = f.read()
            return [Document(page_content=text, metadata={"source": file_path, "file_type": file_type})]

        elif file_type == "docx":
            try:
                doc = docx.Document(file_path)
                text = "\n".join(p.text for p in doc.paragraphs)
                return [Document(page_content=text, metadata={"source": file_path, "file_type": "docx"})]
            except Exception as e:
                logger.error(f"Error loading DOCX file: {e}")
                raise ValueError(f"Failed to load DOCX file: {str(e)}")

        elif file_type == "csv":
            try:
                encodings_to_try = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
                df = None
                for encoding in encodings_to_try:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                if df is None:
                    raise ValueError("Failed to decode CSV with supported encodings")
                documents = []
                row_chunk_size = 20
                for i in range(0, len(df), row_chunk_size):
                    chunk_df = df.iloc[i:i+row_chunk_size]
                    text_block = ""
                    chunk_rows = []
                    for _, row in chunk_df.iterrows():
                        row_dict = row.to_dict()
                        clean_row = {k: v for k, v in row_dict.items() if pd.notna(v)}
                        chunk_rows.append(clean_row)
                        row_text = "\n".join([f"{col}: {val}" for col, val in clean_row.items()])
                        text_block += row_text + "\n\n---\n\n"
                    documents.append(Document(
                        page_content=text_block,
                        metadata={"source": file_path, "file_type": "csv", "row_start": i, "row_data": chunk_rows}
                    ))
                return documents
            except Exception as e:
                logger.error(f"Error loading CSV file: {e}")
                raise ValueError(f"Failed to load CSV file: {str(e)}")

        elif file_type in ["xlsx", "xls"]:
            try:
                df = pd.read_excel(file_path)
                for col in df.select_dtypes(include=['datetime64']).columns:
                    df[col] = df[col].astype(str)
                documents = []
                row_chunk_size = 20
                for i in range(0, len(df), row_chunk_size):
                    chunk_df = df.iloc[i:i+row_chunk_size]
                    text_block = ""
                    chunk_rows = []
                    for _, row in chunk_df.iterrows():
                        row_dict = row.to_dict()
                        clean_row = {k: v for k, v in row_dict.items() if pd.notna(v)}
                        chunk_rows.append(clean_row)
                        row_text = "\n".join([f"{col}: {val}" for col, val in clean_row.items()])
                        text_block += row_text + "\n\n---\n\n"
                    documents.append(Document(
                        page_content=text_block,
                        metadata={"source": file_path, "file_type": file_type, "row_start": i, "row_data": chunk_rows}
                    ))
                return documents
            except Exception as e:
                logger.error(f"Error loading Excel file: {e}")
                raise ValueError(f"Failed to load Excel file: {str(e)}")

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into smaller chunks with section-aware splitting."""
        final_chunks = []
        section_pattern = re.compile(
            r'^(?:\d+\.?\d*\s*)?(Abstract|Introduction|Related Work|Method(?:ology)?|Experiments|Results|Conclusion)\b.*$',
            re.IGNORECASE | re.MULTILINE
        )
        for doc in documents:
            text = doc.page_content
            matches = list(section_pattern.finditer(text))
            blocks = []
            if not matches:
                blocks.append(("General", text))
            else:
                last_idx = 0
                current_section = "General"
                for match in matches:
                    start_idx = match.start()
                    if start_idx > last_idx:
                        blocks.append((current_section, text[last_idx:start_idx].strip()))
                    current_section = match.group(1).title()
                    if current_section == "Methodology":
                        current_section = "Method"
                    last_idx = match.end()
                if last_idx < len(text):
                    blocks.append((current_section, text[last_idx:].strip()))

            for section, block_text in blocks:
                if not block_text:
                    continue
                block_doc = Document(page_content=block_text, metadata={**doc.metadata, "section": section})
                block_chunks = self.text_splitter.split_documents([block_doc])
                final_chunks.extend(block_chunks)

        return final_chunks

    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings using the semaphore-limited embedding service."""
        with _embedding_semaphore:
            return self.embeddings.embed_documents(texts)

    # ─────────────────────────────────────────────────────────────
    # MAIN INGESTION PIPELINE
    # ─────────────────────────────────────────────────────────────
    async def ingest_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Full ingestion pipeline with per-stage diagnostics.
        Logs every failure to ingestion_errors table.
        """
        start_total = time.perf_counter()

        # Diagnostic record built incrementally
        diag = {
            "filename": filename,
            "parse_stage_status": "NOT_STARTED",
            "chunk_count": 0,
            "embedding_stage_status": "NOT_STARTED",
            "db_insert_status": "NOT_STARTED",
        }

        # ── 0. Duplicate check ────────────────────────────────────
        existing = self.db.query(DocumentChunk).filter(
            DocumentChunk.source_file == filename
        ).first()
        if existing:
            logger.warning(f"[DUPLICATE] {filename} already exists — skipping.")
            raise ValueError(f"Document '{filename}' already exists.")

        file_extension = Path(filename).suffix.lower().lstrip('.')

        # ── 1. PARSE ─────────────────────────────────────────────
        start_read = time.perf_counter()
        try:
            documents = self.load_document(file_path, file_extension)
            file_read_ms = (time.perf_counter() - start_read) * 1000

            if not documents:
                diag["parse_stage_status"] = "EMPTY"
                logger.warning(f"[PARSE EMPTY] {filename} produced no documents.")
                return {"filename": filename, "status": "empty", "num_chunks": 0, "diagnostics": diag}

            full_text = "\n".join(d.page_content for d in documents)
            text_len = len(full_text)

            # Phase 3: Validate extracted text
            if text_len < MIN_TEXT_LENGTH:
                msg = f"PDF parsing produced insufficient text ({text_len} chars, minimum {MIN_TEXT_LENGTH})"
                logger.warning(f"[PARSE SHORT] {filename}: {msg}")
                _log_ingestion_error(self.db, filename, "PARSE", ValueError(msg))
                diag["parse_stage_status"] = f"SHORT_TEXT:{text_len}"
                # Proceed anyway — don't abort, log and continue
            else:
                diag["parse_stage_status"] = f"OK:{len(documents)}pages:{text_len}chars"
                logger.info(f"[PARSE OK] {filename} — {len(documents)} pages, {text_len} chars in {file_read_ms:.0f}ms")

        except Exception as e:
            diag["parse_stage_status"] = f"FAILED:{e}"
            _log_ingestion_error(self.db, filename, "PARSE", e)
            logger.error(f"[PARSE FAIL] {filename}: {e}\n{traceback.format_exc()}")
            raise

        # ── Summary extraction (isolated — non-fatal) ──────────
        paper_id = str(uuid.uuid4())
        if file_extension in ["pdf", "txt", "md", "docx"]:
            self._extract_and_store_summary(full_text, filename, paper_id)

        # ── 2. CHUNK ─────────────────────────────────────────────
        start_chunk = time.perf_counter()
        for d in documents:
            d.metadata["paper_id"] = paper_id
        chunks = self.split_documents(documents)
        chunking_ms = (time.perf_counter() - start_chunk) * 1000
        diag["chunk_count"] = len(chunks)
        logger.info(f"[CHUNK OK] {filename} — {len(chunks)} chunks in {chunking_ms:.0f}ms")

        # ── 3. EMBED (semaphore-limited, retries in service) ─────
        texts = [chunk.page_content for chunk in chunks]
        start_embed = time.perf_counter()
        try:
            logger.info(f"[EMBED START] {filename} — {len(texts)} chunks (semaphore slot acquired)")
            with _embedding_semaphore:
                embeddings = self.embeddings.embed_documents(texts)
            embedding_ms = (time.perf_counter() - start_embed) * 1000
            diag["embedding_stage_status"] = f"OK:{len(embeddings)}vectors:{embedding_ms:.0f}ms"
            logger.info(f"[EMBED OK] {filename} — {len(embeddings)} embeddings in {embedding_ms:.0f}ms")
        except Exception as e:
            embedding_ms = (time.perf_counter() - start_embed) * 1000
            diag["embedding_stage_status"] = f"FAILED:{e}"
            _log_ingestion_error(self.db, filename, "EMBEDDING", e)
            logger.error(f"[EMBED FAIL] {filename} after {embedding_ms:.0f}ms: {e}\n{traceback.format_exc()}")
            raise

        # ── 4. DB INSERT (per-chunk safety) ──────────────────────
        start_db = time.perf_counter()
        num_chunks = self._store_chunks_safe(chunks, embeddings, filename, paper_id)
        db_insert_ms = (time.perf_counter() - start_db) * 1000
        diag["db_insert_status"] = f"OK:{num_chunks}rows:{db_insert_ms:.0f}ms"
        logger.info(f"[DB COMMIT OK] {filename} — {num_chunks} chunks in {db_insert_ms:.0f}ms")

        total_ms = (time.perf_counter() - start_total) * 1000
        metrics = {
            "file_read_ms": round(file_read_ms, 2),
            "chunking_ms": round(chunking_ms, 2),
            "embedding_ms": round(embedding_ms, 2),
            "db_insert_ms": round(db_insert_ms, 2),
            "total_ms": round(total_ms, 2)
        }
        logger.info(f"[INGEST COMPLETE] {filename} — {metrics}")

        return {
            "filename": filename,
            "num_chunks": num_chunks,
            "num_pages": len(documents),
            "status": "success",
            "metrics": metrics,
            "diagnostics": diag
        }

    def ingest_document_sync(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Synchronous wrapper for background workers."""
        import asyncio
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        return loop.run_until_complete(self.ingest_document(file_path, filename))

    # ─────────────────────────────────────────────────────────────
    # Phase 4: Per-chunk insert — skip bad chunks, commit per-doc
    # ─────────────────────────────────────────────────────────────
    def _store_chunks_safe(
        self,
        chunks: List[Document],
        embeddings: List[List[float]],
        source_filename: str,
        paper_id: str
    ) -> int:
        """
        Insert chunks one-by-one. If a single chunk fails (e.g. unique violation),
        log and skip it — do NOT abort the entire document.
        Commits once at the end of the document.
        """
        inserted = 0
        skipped = 0

        for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            try:
                metadata = chunk.metadata.copy()
                metadata["original_filename"] = source_filename
                db_chunk = DocumentChunk(
                    source_file=source_filename,
                    chunk_index=idx,
                    chunk_text=chunk.page_content,
                    embedding=embedding,
                    chunk_metadata=metadata
                )
                with self.db.begin_nested():   # SAVEPOINT per chunk
                    self.db.add(db_chunk)
                inserted += 1
            except Exception as chunk_err:
                skipped += 1
                logger.warning(
                    f"[CHUNK SKIP] {source_filename} chunk {idx}: {chunk_err}"
                )
                _log_ingestion_error(
                    self.db, source_filename, f"DB_INSERT_chunk_{idx}", chunk_err
                )

        # Single commit for the whole document
        try:
            self.db.commit()
        except Exception as commit_err:
            self.db.rollback()
            _log_ingestion_error(self.db, source_filename, "DB_COMMIT", commit_err)
            logger.error(f"[DB COMMIT FAIL] {source_filename}: {commit_err}")
            raise

        if skipped:
            logger.warning(f"[DB INSERT] {source_filename} — {inserted} ok, {skipped} skipped")

        # Background graph extraction after commit
        if self.graph_extractor and inserted > 0:
            threading.Thread(
                target=self._background_graph_extract,
                args=(chunks, source_filename),
                daemon=True
            ).start()

        return inserted

    # ─────────────────────────────────────────────────────────────
    # Legacy store_chunks (kept for compatibility, delegates to safe version)
    # ─────────────────────────────────────────────────────────────
    def store_chunks(
        self,
        chunks: List[Document],
        embeddings: List[List[float]],
        source_filename: str
    ) -> int:
        return self._store_chunks_safe(chunks, embeddings, source_filename, "")

    def _extract_and_store_summary(self, text: str, filename: str, paper_id: str):
        """Extract structured summary using LLM — isolated via SAVEPOINT."""
        try:
            with self.db.begin_nested():
                prompt = f"""Extract the following structured information from this research paper:

- Problem statement
- Core contribution
- Methodology summary
- Datasets used
- Evaluation metrics
- Key experimental results
- Limitations
- Main contributions

Return valid JSON only. Format:
{{
  "problem_statement": "...",
  "contributions": "...",
  "methodology": "...",
  "datasets": "...",
  "evaluation_metrics": "...",
  "key_results": "...",
  "limitations": "..."
}}

Paper text (truncate if needed):
{text[:40000]}
"""
                gen_service = GenerationService()
                logger.info(f"Extracting structured summary for {filename}...")
                response = gen_service.generate(prompt)
                match = re.search(r'\{.*\}', response.replace('\n', ' '), re.DOTALL)
                if match:
                    data = json.loads(match.group(0))
                    summary = PaperSummary(
                        id=paper_id,
                        source_file=filename,
                        problem_statement=data.get("problem_statement"),
                        methodology=data.get("methodology"),
                        datasets=data.get("datasets"),
                        evaluation_metrics=data.get("evaluation_metrics"),
                        key_results=data.get("key_results"),
                        limitations=data.get("limitations"),
                        contributions=data.get("contributions")
                    )
                    self.db.add(summary)
                    logger.info(f"Successfully stored PaperSummary for {filename}")
                else:
                    logger.warning(f"Could not parse JSON for summary of {filename}")
        except Exception as e:
            logger.error(f"Summary extraction failed for {filename} (non-fatal): {e}")

    def _background_graph_extract(self, chunks: List[Document], source_filename: str):
        """Run graph extraction in the background."""
        if not self.graph_extractor:
            return
        logger.info(f"Starting background graph extraction for {source_filename}")
        combined_text = ""
        for idx, chunk in enumerate(chunks):
            if idx >= 10:
                break
            combined_text += f"\n--- Section {idx+1} ---\n{chunk.page_content}\n"
        if combined_text:
            try:
                time.sleep(2.0)
                self.graph_extractor.extract_and_store(
                    text=combined_text[:35000],
                    source_file=source_filename,
                    chunk_index=0
                )
            except Exception as graph_err:
                logger.error(f"Background graph extraction failed for {source_filename}: {graph_err}")
        logger.info(f"Background graph extraction completed for {source_filename}")
